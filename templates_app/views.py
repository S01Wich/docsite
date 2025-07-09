import re
import os
from datetime import datetime

from django.shortcuts import render, get_object_or_404
from django.http import FileResponse
from django.conf import settings

from docx import Document
from docx.oxml.ns import qn

from .models import TemplateDoc
from .forms import generate_template_form

# Паттерн: только {$Тег}, без учёта стилей внутри
TAG_PATTERN = r'\{\$(?P<tag>[A-Za-zА-Яа-я0-9_]+)\}'

def index(request):
    templates = TemplateDoc.objects.all()
    return render(request, 'index.html', {'templates': templates})

def fill_template(request, pk):
    template = get_object_or_404(TemplateDoc, pk=pk)
    doc_path = template.doc_file.path

    # Загружаем документ и собираем все теги {$...}
    doc = Document(doc_path)
    tags = set()
    for p in doc.paragraphs:
        for m in re.finditer(TAG_PATTERN, p.text):
            tags.add(m.group('tag'))
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for m in re.finditer(TAG_PATTERN, p.text):
                        tags.add(m.group('tag'))

    # Новая сортировка:
    # 1) все теги, начинающиеся с 'ФИО' и не содержащие числового суффикса — первыми
    # 2) затем по группам с числовым суффиксом '_<номер>':
    #    - сначала все теги, которые начинаются с 'ФИО' и заканчиваются на '_<номер>'
    #    - затем все остальные теги с тем же суффиксом (по алфавиту)
    # 3) в конце — все прочие теги без числового суффикса (по алфавиту)
    def sort_key(tag):
        # Проверяем числовой суффикс
        m_num = re.match(r'^(.*)_(\d+)$', tag)
        if m_num:
            base = m_num.group(1)
            num = int(m_num.group(2))
            # Внутри группы num: теги, начинающиеся с 'ФИО', первыми
            if base.startswith('ФИО'):
                return (num, 0, base.lower())
            return (num, 1, base.lower())
        # Без числового суффикса
        if tag.startswith('ФИО'):
            # Все теги, начинающиеся с 'ФИО', без номера — перед всеми группами
            return (-1, 0, tag.lower())
        # Прочие теги без номера — после всех
        return (float('inf'), 1, tag.lower())

    sorted_tags = sorted(tags, key=sort_key)

    # Формируем форму с полями по отсортированным тегам
    TemplateForm = generate_template_form(sorted_tags)

    if request.method == 'POST':
        form = TemplateForm(request.POST)
        if form.is_valid():
            # Создаём копию документа для правок
            output_doc = Document(doc_path)

            # Заменяем теги в параграфах, сохраняя стили
            for paragraph in output_doc.paragraphs:
                for run in paragraph.runs:
                    if re.search(TAG_PATTERN, run.text):
                        run.text = re.sub(
                            TAG_PATTERN,
                            lambda m: str(form.cleaned_data.get(m.group('tag'), '')),
                            run.text
                        )
                    run.font.name = 'Times New Roman'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

            # Заменяем теги в таблицах, сохраняя стили
            for table in output_doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                if re.search(TAG_PATTERN, run.text):
                                    run.text = re.sub(
                                        TAG_PATTERN,
                                        lambda m: str(form.cleaned_data.get(m.group('tag'), '')),
                                        run.text
                                    )
                                run.font.name = 'Times New Roman'
                                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

            # Формируем имя выходного файла: <оригинал>_<YYYY-MM-DD>.docx
            basename = os.path.basename(doc_path)
            name, ext = os.path.splitext(basename)
            if not ext:
                ext = '.docx'
            date_str = datetime.now().strftime('%Y-%m-%d')
            output_filename = f"{name}_{date_str}{ext}"

            # Сохраняем и возвращаем файл
            os.makedirs(settings.MEDIA_ROOT, exist_ok=True)
            output_path = os.path.join(settings.MEDIA_ROOT, output_filename)
            output_doc.save(output_path)

            return FileResponse(
                open(output_path, 'rb'),
                as_attachment=True,
                filename=output_filename
            )
    else:
        form = TemplateForm()

    return render(request, 'fill_template.html', {
        'template': template,
        'form': form,
    })
